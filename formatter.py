"""
Professional .docx question paper formatter.
Optimized for MINIMAL paper usage while maintaining readability.
Uses python-docx for document generation.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
from datetime import datetime


def set_cell_border(cell, **kwargs):
    """Set cell border. Usage: set_cell_border(cell, top={"sz":4, "color":"000000"})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" w:sz="{attrs.get("sz", 4)}" '
            f'w:space="0" w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def create_question_paper(
    structured_data: dict,
    output_path: str,
    school_name: str = "",
    logo_path: str = None,
    compact: bool = True,
    question_images: dict = None
) -> str:
    """
    Generate a professional .docx question paper from structured data.
    
    Args:
        structured_data: Dict with exam_title, sections, questions etc.
        output_path: Where to save the .docx file
        school_name: School name for header
        logo_path: Path to school logo image
        compact: If True, optimize for minimal paper usage
    
    Returns:
        Path to the generated .docx file
    """
    doc = Document()
    
    # ─── Page Setup: Compact margins ───────────────────────────────────────
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)    # A4
        section.page_height = Cm(29.7)
        if compact:
            section.top_margin = Cm(1.2)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
        else:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)

    # ─── Default font setup ────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11) if compact else Pt(12)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(1) if compact else Pt(3)
    style.paragraph_format.line_spacing = 1.0 if compact else 1.15

    data = structured_data
    
    # ─── HEADER SECTION ────────────────────────────────────────────────────
    
    # School Logo + Name (using table for side-by-side layout)
    display_school = school_name or data.get("school_name", "")
    
    if logo_path and os.path.exists(logo_path) and display_school:
        # Logo + School Name side by side
        header_table = doc.add_table(rows=1, cols=2)
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Logo cell
        logo_cell = header_table.cell(0, 0)
        logo_cell.width = Cm(2.5)
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = logo_para.add_run()
        run.add_picture(logo_path, height=Cm(1.8))
        
        # School name cell
        name_cell = header_table.cell(0, 1)
        name_para = name_cell.paragraphs[0]
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = name_para.add_run(display_school.upper())
        run.bold = True
        run.font.size = Pt(14) if compact else Pt(16)
        run.font.name = 'Times New Roman'
        
        # Remove table borders
        for row in header_table.rows:
            for cell in row.cells:
                set_cell_border(cell,
                    top={"sz": 0, "color": "FFFFFF"},
                    bottom={"sz": 0, "color": "FFFFFF"},
                    left={"sz": 0, "color": "FFFFFF"},
                    right={"sz": 0, "color": "FFFFFF"})
    elif display_school:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(display_school.upper())
        run.bold = True
        run.font.size = Pt(14) if compact else Pt(16)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(0)
    elif logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, height=Cm(2.0))
        p.paragraph_format.space_after = Pt(0)

    # Exam Title
    exam_title = data.get("exam_title", "")
    if exam_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(exam_title)
        run.bold = True
        run.font.size = Pt(12) if compact else Pt(13)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    # ─── Metadata line (Class | Subject | Time | Marks) - single line ─────
    meta_parts = []
    if data.get("class"):
        meta_parts.append(f"Class: {data['class']}")
    if data.get("subject"):
        meta_parts.append(f"Subject: {data['subject']}")
    if data.get("time"):
        meta_parts.append(f"Time: {data['time']}")
    if data.get("total_marks"):
        meta_parts.append(f"Max. Marks: {data['total_marks']}")

    if meta_parts:
        # Use a table for clean alignment: left side and right side
        if len(meta_parts) >= 4:
            meta_table = doc.add_table(rows=2, cols=2)
            meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Row 1: Class (left) | Time (right)
            left1 = meta_table.cell(0, 0).paragraphs[0]
            run = left1.add_run(meta_parts[0])
            run.font.size = Pt(10) if compact else Pt(11)
            run.font.name = 'Times New Roman'
            left1.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            right1 = meta_table.cell(0, 1).paragraphs[0]
            run = right1.add_run(meta_parts[2])
            run.font.size = Pt(10) if compact else Pt(11)
            run.font.name = 'Times New Roman'
            right1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Row 2: Subject (left) | Marks (right)
            left2 = meta_table.cell(1, 0).paragraphs[0]
            run = left2.add_run(meta_parts[1])
            run.font.size = Pt(10) if compact else Pt(11)
            run.font.name = 'Times New Roman'
            left2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            right2 = meta_table.cell(1, 1).paragraphs[0]
            run = right2.add_run(meta_parts[3])
            run.font.size = Pt(10) if compact else Pt(11)
            run.font.name = 'Times New Roman'
            right2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Remove borders
            for row in meta_table.rows:
                for cell in row.cells:
                    set_cell_border(cell,
                        top={"sz": 0, "color": "FFFFFF"},
                        bottom={"sz": 0, "color": "FFFFFF"},
                        left={"sz": 0, "color": "FFFFFF"},
                        right={"sz": 0, "color": "FFFFFF"})
                    for para in cell.paragraphs:
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("  |  ".join(meta_parts))
            run.font.size = Pt(10) if compact else Pt(11)
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_after = Pt(2)

    # ─── Divider line ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ─── Instructions ──────────────────────────────────────────────────────
    instructions = data.get("instructions", [])
    if instructions:
        p = doc.add_paragraph()
        run = p.add_run("General Instructions:")
        run.bold = True
        run.font.size = Pt(10) if compact else Pt(11)
        run.font.name = 'Times New Roman'
        run.underline = True
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)

        for idx, instr in enumerate(instructions, 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            
            run = p.add_run(f"{idx}. ")
            run.font.size = Pt(9) if compact else Pt(10)
            run.font.name = 'Times New Roman'
            
            run = p.add_run(instr)
            run.font.size = Pt(9) if compact else Pt(10)
            run.font.name = 'Times New Roman'

    # ─── Another divider ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ─── SECTIONS & QUESTIONS ──────────────────────────────────────────────
    sections = data.get("sections", [])
    
    for si, section in enumerate(sections):
        section_name = section.get("section_name", f"Section {si + 1}")
        
        # Section header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(section_name.upper())
        run.bold = True
        run.font.size = Pt(11) if compact else Pt(12)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(6) if compact else Pt(10)
        p.paragraph_format.space_after = Pt(3) if compact else Pt(6)

        questions = section.get("questions", [])
        
        for qi, question in enumerate(questions):
            q_num = question.get("number", str(qi + 1))
            q_text = question.get("text", "")
            q_marks = question.get("marks", "")
            subparts = question.get("subparts", [])
            
            # ── Question with marks on the right using tab stop ──
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3) if compact else Pt(4)
            p.paragraph_format.space_after = Pt(1) if compact else Pt(2)
            p.paragraph_format.line_spacing = 1.0 if compact else 1.1
            
            # Add tab stop for right-aligned marks
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(18.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            
            # Question number (bold)
            run = p.add_run(f"Q{q_num}. ")
            run.bold = True
            run.font.size = Pt(10.5) if compact else Pt(11)
            run.font.name = 'Times New Roman'
            
            # Question text
            run = p.add_run(q_text)
            run.font.size = Pt(10.5) if compact else Pt(11)
            run.font.name = 'Times New Roman'
            
            # Marks (right-aligned via tab)
            if q_marks:
                run = p.add_run(f"\t[{q_marks}]")
                run.bold = True
                run.font.size = Pt(10) if compact else Pt(11)
                run.font.name = 'Times New Roman'
            
            # ── Subparts ──
            if subparts:
                # Detect subpart type
                is_mcq = (len(subparts) == 4 and 
                          all(sp.strip().startswith(('(a)', '(b)', '(c)', '(d)', 
                                                     'a)', 'b)', 'c)', 'd)',
                                                     'A)', 'B)', 'C)', 'D)',
                                                     '(A)', '(B)', '(C)', '(D)')) 
                                for sp in subparts))
                
                # Detect match-the-following / two-column data
                # Look for tab characters, multiple spaces (3+), or arrow/dash separators
                import re as _re
                is_match_columns = any(
                    ('\t' in sp or '  ' in sp.strip() or 
                     _re.search(r'\s{3,}', sp) or
                     ' → ' in sp or ' -> ' in sp or ' – ' in sp or ' — ' in sp)
                    for sp in subparts
                )
                
                if is_match_columns:
                    # ── Match-the-following: render as a 2-column table ──
                    # Parse each line into two columns
                    rows_data = []
                    for sp in subparts:
                        sp = sp.strip()
                        # Try splitting by tab, arrow, or multiple spaces
                        parts = None
                        if '\t' in sp:
                            parts = [x.strip() for x in sp.split('\t', 1)]
                        elif ' → ' in sp:
                            parts = [x.strip() for x in sp.split(' → ', 1)]
                        elif ' -> ' in sp:
                            parts = [x.strip() for x in sp.split(' -> ', 1)]
                        elif ' — ' in sp:
                            parts = [x.strip() for x in sp.split(' — ', 1)]
                        elif ' – ' in sp:
                            parts = [x.strip() for x in sp.split(' – ', 1)]
                        elif _re.search(r'\s{3,}', sp):
                            parts = [x.strip() for x in _re.split(r'\s{3,}', sp, maxsplit=1)]
                        
                        if parts and len(parts) == 2:
                            rows_data.append(parts)
                        else:
                            rows_data.append([sp, ""])
                    
                    match_table = doc.add_table(rows=len(rows_data), cols=2)
                    match_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    
                    for ri, (col_a, col_b) in enumerate(rows_data):
                        for ci, text in enumerate([col_a, col_b]):
                            cell = match_table.cell(ri, ci)
                            cell_para = cell.paragraphs[0]
                            cell_para.paragraph_format.space_before = Pt(1)
                            cell_para.paragraph_format.space_after = Pt(1)
                            cell_para.paragraph_format.left_indent = Cm(0.2)
                            run = cell_para.add_run(text)
                            run.font.size = Pt(10) if compact else Pt(11)
                            run.font.name = 'Times New Roman'
                            
                            # Light borders for match tables
                            set_cell_border(cell,
                                top={"sz": 4, "color": "CCCCCC"},
                                bottom={"sz": 4, "color": "CCCCCC"},
                                left={"sz": 4, "color": "CCCCCC"},
                                right={"sz": 4, "color": "CCCCCC"})
                    
                    for row in match_table.rows:
                        tr = row._tr
                        trPr = tr.get_or_add_trPr()
                        trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="300" w:hRule="atLeast"/>')
                        trPr.append(trHeight)
                
                elif is_mcq and compact:
                    # ── MCQ: 2x2 grid ──
                    opt_table = doc.add_table(rows=2, cols=2)
                    opt_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    
                    for oi, opt in enumerate(subparts):
                        row_idx = oi // 2
                        col_idx = oi % 2
                        cell = opt_table.cell(row_idx, col_idx)
                        cell_para = cell.paragraphs[0]
                        cell_para.paragraph_format.space_before = Pt(0)
                        cell_para.paragraph_format.space_after = Pt(0)
                        cell_para.paragraph_format.left_indent = Cm(0.3)
                        run = cell_para.add_run(opt.strip())
                        run.font.size = Pt(10) if compact else Pt(11)
                        run.font.name = 'Times New Roman'
                        
                        set_cell_border(cell,
                            top={"sz": 0, "color": "FFFFFF"},
                            bottom={"sz": 0, "color": "FFFFFF"},
                            left={"sz": 0, "color": "FFFFFF"},
                            right={"sz": 0, "color": "FFFFFF"})
                    
                    for row in opt_table.rows:
                        tr = row._tr
                        trPr = tr.get_or_add_trPr()
                        trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="280" w:hRule="atLeast"/>')
                        trPr.append(trHeight)
                else:
                    # ── Regular subparts ──
                    for sp in subparts:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(1.2)
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.0
                        
                        run = p.add_run(sp.strip())
                        run.font.size = Pt(10) if compact else Pt(11)
                        run.font.name = 'Times New Roman'

            # ── Question Image ──
            if question_images:
                img_key = f"{si}_{qi}"
                if img_key in question_images:
                    img_path = question_images[img_key]
                    if os.path.exists(img_path):
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Cm(0.5)
                        p.paragraph_format.space_before = Pt(4)
                        p.paragraph_format.space_after = Pt(4)
                        p.paragraph_format.line_spacing = 1.0
                        
                        # Calculate max width based on mode
                        max_width_cm = 8.0 if compact else 10.0
                        
                        try:
                            from PIL import Image as PILImage
                            with PILImage.open(img_path) as img:
                                w, h = img.size
                                aspect = h / w
                                width_cm = min(max_width_cm, w * 0.0264583)  # px to cm approx
                                height_cm = width_cm * aspect
                                # Cap height to avoid full-page images
                                max_height_cm = 8.0 if compact else 10.0
                                if height_cm > max_height_cm:
                                    height_cm = max_height_cm
                                    width_cm = height_cm / aspect
                                run = p.add_run()
                                run.add_picture(img_path, width=Cm(width_cm))
                        except ImportError:
                            # No PIL, just insert with fixed width
                            run = p.add_run()
                            run.add_picture(img_path, width=Cm(max_width_cm))

    # ─── Footer: End of Paper ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of Question Paper —")
    run.italic = True
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ─── Page numbers in footer ────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        
        run = fp.add_run("Page ")
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Add page number field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run1 = fp.add_run()
        run1._r.append(fldChar1)
        
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2 = fp.add_run()
        run2._r.append(instrText)
        run2.font.size = Pt(8)
        run2.font.name = 'Times New Roman'
        run2.font.color.rgb = RGBColor(128, 128, 128)
        
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3 = fp.add_run()
        run3._r.append(fldChar2)

    # ─── Save ──────────────────────────────────────────────────────────────
    doc.save(output_path)
    return output_path


def generate_filename(data: dict) -> str:
    """Generate a descriptive filename from structured data."""
    parts = []
    if data.get("class"):
        parts.append(f"Class_{data['class']}")
    if data.get("subject"):
        parts.append(data['subject'].replace(" ", "_"))
    if data.get("exam_title"):
        # Take first few words
        title_words = data['exam_title'].split()[:3]
        parts.append("_".join(title_words))
    
    parts.append(datetime.now().strftime("%Y%m%d"))
    
    return "_".join(parts) + ".docx" if parts else "Question_Paper.docx"
